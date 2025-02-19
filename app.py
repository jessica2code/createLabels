import os
from flask import Flask, render_template, request, redirect, url_for, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from collections import defaultdict
import pandas as pd
import shutil

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'generated_files'
app.config['ALLOWED_EXTENSIONS'] = {'xls', 'xlsx'}

class LabelGenerator:
    def __init__(self):
        self.data = defaultdict(list)
        self.default_made_in = "MADE IN CHINA"
        self.current_factory = None

    def set_customer_info(self, order_number, customer_name, customer_info, made_in,
                          name_font_size, product_info_font_size, product_quantity_font_size,
                          made_in_font_size, customer_info_font_size):
        self.order_number = order_number
        self.customer_name = customer_name
        self.customer_info = customer_info
        self.made_in = made_in if made_in else self.default_made_in
        self.name_font_size = Pt(int(name_font_size)) if name_font_size else Pt(12)
        self.product_info_font_size = Pt(int(product_info_font_size)) if product_info_font_size else Pt(12)
        self.product_quantity_font_size = Pt(int(product_quantity_font_size)) if product_quantity_font_size else Pt(12)
        self.made_in_font_size = Pt(int(made_in_font_size)) if made_in_font_size else Pt(12)
        self.customer_info_font_size = Pt(int(customer_info_font_size)) if customer_info_font_size else Pt(12)

    def set_current_factory(self, factory_name):
        self.current_factory = factory_name

    def add_product(self, product_info, product_quantity):
        if self.current_factory:
            self.data[self.current_factory].append((self.customer_name, product_info, product_quantity, self.made_in, self.customer_info))

    def add_products_from_file(self, file_path):
        df = pd.read_excel(file_path)
        for index, row in df.iterrows():
            factory_name = row['工厂名']
            product_info = row['产品信息']
            product_quantity = row['产品数量']
            self.data[factory_name].append((self.customer_name, product_info, product_quantity, self.made_in, self.customer_info))

    def generate_documents(self):
        folder_name = os.path.join(app.config['UPLOAD_FOLDER'], self.order_number)
        if not os.path.exists(folder_name):
            os.makedirs(folder_name)
        for factory_name, products in self.data.items():
            doc = Document()
            for i, product in enumerate(products):
                customer_name, product_info, product_quantity, made_in, customer_info = product
                p1 = doc.add_paragraph()
                run1 = p1.add_run(customer_name)
                run1.font.size = self.name_font_size
                
                p2 = doc.add_paragraph()
                run2 = p2.add_run(f"ITEM: {product_info}")
                run2.font.size = self.product_info_font_size
                
                p3 = doc.add_paragraph()
                run3 = p3.add_run(f"QTY: {product_quantity}")
                run3.font.size = self.product_quantity_font_size
                
                p4 = doc.add_paragraph()
                run4 = p4.add_run(made_in)
                run4.font.size = self.made_in_font_size
                
                p5 = doc.add_paragraph()
                run5 = p5.add_run(customer_info)
                run5.font.size = self.customer_info_font_size
                
                if i < len(products) - 1:  # 在每个产品信息后面添加隔页符，最后一个产品除外
                    doc.add_page_break()
            doc.save(os.path.join(folder_name, f"{factory_name}.docx"))
        return folder_name

    def create_zip(self, folder_name):
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{self.order_number}.zip")
        shutil.make_archive(zip_path.replace('.zip', ''), 'zip', folder_name)
        return zip_path

label_generator = LabelGenerator()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        order_number = request.form['order_number']
        customer_name = request.form['customer_name']
        customer_info = request.form['customer_info']
        made_in = request.form['made_in']
        name_font_size = request.form['name_font_size']
        product_info_font_size = request.form['product_info_font_size']
        product_quantity_font_size = request.form['product_quantity_font_size']
        made_in_font_size = request.form['made_in_font_size']
        customer_info_font_size = request.form['customer_info_font_size']
        
        label_generator.set_customer_info(order_number, customer_name, customer_info, made_in, 
                                          name_font_size, product_info_font_size, product_quantity_font_size, 
                                          made_in_font_size, customer_info_font_size)

        if 'file' in request.files:
            file = request.files['file']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                label_generator.add_products_from_file(file_path)
                folder_name = label_generator.generate_documents()
                zip_path = label_generator.create_zip(folder_name)
                return send_file(zip_path, as_attachment=True)
    return render_template('index.html')

@app.route('/return_home')
def return_home():
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)